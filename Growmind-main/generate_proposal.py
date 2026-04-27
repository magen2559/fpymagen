from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_proposal():
    doc = Document()
    
    # Title
    title = doc.add_heading('GROW MIND: Cinematic Productivity Ecosystem', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('FYP Proposal Upgrade', level=1)

    # 1. Abstract
    doc.add_heading('1. Abstract', level=2)
    doc.add_paragraph(
        'In the modern digital landscape, university students face significant challenges in maintaining academic focus. '
        'This project introduces GrowMind, an immersive, WebGL-driven productivity ecosystem designed to cultivate "deep work" habits. '
        'Unlike traditional focus apps, GrowMind utilizes React Three Fiber and Custom GLSL Shaders to transform study hours into the organic growth of a cinematic virtual garden. '
        'By synthesizing behavioral psychology with a premium high-performance tech stack, GrowMind provides immediate visual gratification and intelligent AI coaching to reclaim student academic potential.'
    )

    # 2. Problem Statement
    doc.add_heading('2. Problem Statement', level=2)
    p = doc.add_paragraph()
    p.add_run('Current productivity tools often feel static and utilitarian, lacking the "dopamine-driven" engagement found in modern digital entertainment. This project addresses:').bold = True
    doc.add_paragraph('1. Delayed Gratification: Academic success lacks immediate rewards.', style='List Bullet')
    doc.add_paragraph('2. Digital Fragmentation: Constant context-switching between study and entertainment.', style='List Bullet')
    doc.add_paragraph('3. Academic Isolation: The lack of community presence during solitary study sessions.', style='List Bullet')

    # 3. Project Aims and Objectives
    doc.add_heading('3. Project Aims and Objectives', level=2)
    doc.add_heading('Aims', level=3)
    doc.add_paragraph('To engineer a sustainable "deep work" habit using a high-fidelity, immediate gratification loop through WebGL immersion and Realtime community visualization.')
    
    doc.add_heading('Objectives', level=3)
    doc.add_paragraph('Design a 60FPS WebGL virtual greenhouse using React Native and Expo.', style='List Bullet')
    doc.add_paragraph('Implement organic growth mechanics using procedural GLSL shaders.', style='List Bullet')
    doc.add_paragraph('Integrate Google Gemini AI as a point-cloud-rendered academic companion.', style='List Bullet')
    doc.add_paragraph('Visualize a "Global Study Garden" using Supabase Realtime to mitigate isolation.', style='List Bullet')

    # 4. Proposed Technical Stack
    doc.add_heading('4. Proposed Technical Stack', level=2)
    doc.add_paragraph('Frontend: React Native + Expo (SDK 52+)', style='List Bullet')
    doc.add_paragraph('3D/Graphics: React Three Fiber (R3F) via expo-gl, GSAP', style='List Bullet')
    doc.add_paragraph('Backend: Supabase (PostgreSQL, Realtime, Edge Functions)', style='List Bullet')
    doc.add_paragraph('Backend Logic: TypeScript / Supabase Edge Functions', style='List Bullet')
    doc.add_paragraph('Styling: Native Components with HSL design tokens and GLSL Shaders', style='List Bullet')

    # 5. Technical Specifications
    doc.add_heading('5. Technical Specifications', level=2)
    
    specs = [
        ('Platform Architecture', 'Native Mobile App (React Native + Expo SDK 52)'),
        ('Core Engine', 'React Three Fiber + Custom GLSL Shaders for organic growth rendering'),
        ('Motion Orchestration', 'GSAP & ScrollTrigger for premium UI sequences'),
        ('Data & Auth', 'Supabase PostgreSQL + Supabase Auth'),
        ('Logic & Services', 'Supabase Edge Functions (TypeScript/Deno)'),
        ('Realtime Services', 'Supabase Realtime for community "Wisp" visualization')
    ]

    for label, desc in specs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{label}: ').bold = True
        p.add_run(desc)

    # 6. Implementation Timeline
    doc.add_heading('6. Implementation Timeline (12 Weeks)', level=2)
    doc.add_paragraph('Weeks 1-4: Core WebGL Engine & Design System.', style='List Bullet')
    doc.add_paragraph('Weeks 5-8: Realtime Integration & Shaders Growth System.', style='List Bullet')
    doc.add_paragraph('Weeks 9-12: AI Companion, Polish (Post-FX), and Documentation.', style='List Bullet')

    doc.save('GrowMind_Final_Proposal.docx')

if __name__ == "__main__":
    create_proposal()
    print("Proposal generated: GrowMind_Final_Proposal.docx")
